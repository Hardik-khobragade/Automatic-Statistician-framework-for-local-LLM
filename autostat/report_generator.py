"""LAYER 4: Report Generator.

Assembles a polished .docx report from:
  - the data profile (Layer 1)
  - the agent's narrative report_sections, with their attached figures (Layer 2)
  - every dict passed to record_result() during the run (Layer 3) -- rendered
    into tables by render_result_dict() below, independent of the LLM's
    prose, so the numbers in the report are exactly what the toolkit
    computed even if the model's narrative is mediocre.
  - validation warnings and an optional ground-truth comparison (Layer 5)

Also generates a baseline "Data Overview" page (missingness chart, a small
multiples histogram grid, a correlation heatmap) itself -- deterministically,
not via the LLM -- so the report looks complete even if a small local model
didn't get around to visualizing much.
"""
from __future__ import annotations

import datetime
import os
from typing import Any, Dict, List, Optional

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .react_agent import ReportSection

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
LIGHT_GREY = RGBColor(0x59, 0x59, 0x59)


# --------------------------------------------------------------------------- #
# Document-level setup
# --------------------------------------------------------------------------- #

def _setup_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size, color in [(1, 20, ACCENT), (2, 16, ACCENT), (3, 13, ACCENT)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def _add_page_number_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _add_title_page(document: Document, title: str, subtitle: str, meta_lines: List[str]) -> None:
    for _ in range(3):
        document.add_paragraph()
    run = document.add_paragraph().add_run(title)
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    run2 = document.add_paragraph().add_run(subtitle)
    run2.font.size = Pt(14)
    run2.font.color.rgb = LIGHT_GREY
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    for line in meta_lines:
        mp = document.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = mp.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = LIGHT_GREY
    document.add_page_break()


# --------------------------------------------------------------------------- #
# Generic table helpers
# --------------------------------------------------------------------------- #

def _style_table(table) -> None:
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _fmt(v: Any) -> str:
    if v is None:
        return "-"
    if isinstance(v, bool):
        return "Yes" if v else "No"
    if isinstance(v, float):
        return f"{v:.4g}"
    if isinstance(v, dict):
        return "; ".join(f"{k}={_fmt(val)}" for k, val in v.items())
    if isinstance(v, (list, tuple)):
        return ", ".join(_fmt(x) for x in v)
    return str(v)


def _add_kv_table(document: Document, items: List[tuple]) -> None:
    table = document.add_table(rows=0, cols=2)
    _style_table(table)
    for k, v in items:
        row = table.add_row().cells
        row[0].text = str(k).replace("_", " ").capitalize()
        row[1].text = _fmt(v)


def _add_matrix_table(document: Document, title: str, outer: Dict[str, Dict[str, Any]]) -> None:
    inner_keys: List[str] = []
    for v in outer.values():
        for k in v.keys():
            if k not in inner_keys:
                inner_keys.append(k)
    document.add_paragraph(title).runs[0].bold = True
    table = document.add_table(rows=1, cols=1 + len(inner_keys))
    _style_table(table)
    hdr = table.rows[0].cells
    hdr[0].text = ""
    for j, k in enumerate(inner_keys):
        hdr[j + 1].text = str(k).replace("_", " ")
    for outer_key, inner in outer.items():
        cells = table.add_row().cells
        cells[0].text = str(outer_key)
        for j, k in enumerate(inner_keys):
            cells[j + 1].text = _fmt(inner.get(k))


def _add_list_of_dict_table(document: Document, title: str, rows: List[Dict[str, Any]]) -> None:
    if not rows:
        return
    keys = list(rows[0].keys())
    if title:
        document.add_paragraph(title).runs[0].bold = True
    table = document.add_table(rows=1, cols=len(keys))
    _style_table(table)
    for j, k in enumerate(keys):
        table.rows[0].cells[j].text = str(k).replace("_", " ")
    for r in rows:
        cells = table.add_row().cells
        for j, k in enumerate(keys):
            cells[j].text = _fmt(r.get(k))


def render_result_dict(document: Document, result: Dict[str, Any]) -> None:
    """Generic renderer: works for every stats_toolkit result shape without
    per-test-specific code, by classifying each field as scalar / flat dict
    / dict-of-dicts / list-of-dicts and rendering accordingly."""
    name = result.get("name") or result.get("test", "Result")
    document.add_heading(str(name).replace("_", " ").title(), level=3)

    scalars, flat_dicts, matrix_dicts, list_dicts, other = [], [], [], [], []
    for k, v in result.items():
        if k == "name":
            continue
        if isinstance(v, dict):
            if v and all(isinstance(vv, dict) for vv in v.values()):
                matrix_dicts.append((k, v))
            elif v:
                flat_dicts.append((k, v))
        elif isinstance(v, list) and v and all(isinstance(it, dict) for it in v):
            list_dicts.append((k, v))
        elif isinstance(v, (int, float, str, bool)) or v is None:
            scalars.append((k, v))
        elif isinstance(v, list):
            # list of scalars (e.g. group labels, per-group means) -- _fmt joins them
            scalars.append((k, v))
        else:
            other.append((k, v))

    if scalars:
        _add_kv_table(document, scalars)
    for k, v in flat_dicts:
        document.add_paragraph(str(k).replace("_", " ").capitalize()).runs[0].bold = True
        _add_kv_table(document, list(v.items()))
    for k, v in matrix_dicts:
        _add_matrix_table(document, str(k).replace("_", " ").capitalize(), v)
    for k, v in list_dicts:
        _add_list_of_dict_table(document, str(k).replace("_", " ").capitalize(), v)
    for k, v in other:
        document.add_paragraph(f"{k}: {v}")
    document.add_paragraph()


# --------------------------------------------------------------------------- #
# Baseline auto-generated charts (deterministic, not LLM-authored)
# --------------------------------------------------------------------------- #

def _gen_missingness_chart(df: pd.DataFrame, path: str) -> Optional[str]:
    miss = df.isna().mean().sort_values(ascending=False)
    miss = miss[miss > 0]
    if miss.empty:
        return None
    fig, ax = plt.subplots(figsize=(7, max(2, 0.35 * len(miss))))
    ax.barh(miss.index[::-1], (miss.values[::-1] * 100), color="#1F4E79")
    ax.set_xlabel("% missing")
    ax.set_title("Missing data by column")
    fig.tight_layout()
    fig.savefig(path, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return path


def _gen_distribution_grid(df: pd.DataFrame, numeric_cols: List[str], path: str) -> Optional[str]:
    cols = numeric_cols[:6]
    if not cols:
        return None
    n = len(cols)
    ncols = min(3, n)
    nrows = (n + ncols - 1) // ncols
    fig, axes = plt.subplots(nrows, ncols, figsize=(4 * ncols, 3 * nrows))
    axes = np.atleast_1d(axes).flatten()
    for ax, col in zip(axes, cols):
        data = pd.to_numeric(df[col], errors="coerce").dropna()
        ax.hist(data, bins=25, color="#2E75B6", edgecolor="white")
        ax.set_title(col, fontsize=10)
    for ax in axes[len(cols):]:
        ax.axis("off")
    fig.suptitle("Distributions of key numeric variables")
    fig.tight_layout()
    fig.savefig(path, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return path


def _gen_correlation_heatmap(df: pd.DataFrame, numeric_cols: List[str], path: str) -> Optional[str]:
    cols = numeric_cols[:15]
    if len(cols) < 2:
        return None
    corr = df[cols].apply(pd.to_numeric, errors="coerce").corr()
    fig, ax = plt.subplots(figsize=(0.6 * len(cols) + 2, 0.6 * len(cols) + 2))
    im = ax.imshow(corr.values, cmap="RdBu_r", vmin=-1, vmax=1)
    ax.set_xticks(range(len(cols)))
    ax.set_yticks(range(len(cols)))
    ax.set_xticklabels(cols, rotation=45, ha="right", fontsize=8)
    ax.set_yticklabels(cols, fontsize=8)
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    ax.set_title("Correlation heatmap")
    fig.tight_layout()
    fig.savefig(path, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return path


# --------------------------------------------------------------------------- #
# Main entry point
# --------------------------------------------------------------------------- #

def build_report(
    output_path: str,
    dataset_name: str,
    task: str,
    df: pd.DataFrame,
    profile: Dict[str, Any],
    report_sections: List[ReportSection],
    recorded_results: List[Dict[str, Any]],
    figures_dir: str,
    model_name: str = "",
    validation_warnings: Optional[List[str]] = None,
    ground_truth_comparison: Optional[Dict[str, Any]] = None,
    executive_summary: Optional[str] = None,
) -> str:
    document = Document()
    _setup_styles(document)

    meta = [f"Dataset: {dataset_name}", f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M}"]
    if model_name:
        meta.append(f"Analysis model: {model_name}")
    _add_title_page(document, "Automatic Statistician Report", task, meta)
    _add_page_number_footer(document)

    document.add_heading("Table of Contents", level=1)
    _add_toc(document)
    document.add_page_break()

    # --- Executive summary -------------------------------------------------
    document.add_heading("Executive Summary", level=1)
    if executive_summary:
        document.add_paragraph(executive_summary)
    else:
        document.add_paragraph(
            f"This dataset contains {profile['n_rows']} rows and {profile['n_cols']} columns. "
            f"{profile['total_missing_cells']} cells were missing and {profile['duplicate_rows']} "
            f"duplicate rows were found. {len(recorded_results)} statistical test(s) were run and "
            f"{len(report_sections)} narrative finding(s) were recorded by the analysis agent; see the "
            f"sections below for details."
        )

    # --- Data overview (auto-generated, deterministic) --------------------
    document.add_heading("Data Overview", level=1)
    cols_info = profile["columns"]
    type_counts: Dict[str, int] = {}
    for info in cols_info.values():
        type_counts[info["detected_type"]] = type_counts.get(info["detected_type"], 0) + 1
    _add_kv_table(document, [
        ("Rows", profile["n_rows"]), ("Columns", profile["n_cols"]),
        ("Missing cells", profile["total_missing_cells"]), ("Duplicate rows", profile["duplicate_rows"]),
    ] + [(f"{t} columns", c) for t, c in type_counts.items()])

    numeric_cols = [c for c, i in cols_info.items() if i["detected_type"] in ("numeric", "ordinal_candidate")]

    miss_path = _gen_missingness_chart(df, os.path.join(figures_dir, "_auto_missingness.png"))
    if miss_path:
        document.add_paragraph()
        document.add_picture(miss_path, width=Inches(6))

    dist_path = _gen_distribution_grid(df, numeric_cols, os.path.join(figures_dir, "_auto_distributions.png"))
    if dist_path:
        document.add_picture(dist_path, width=Inches(6))

    corr_path = _gen_correlation_heatmap(df, numeric_cols, os.path.join(figures_dir, "_auto_corr.png"))
    if corr_path:
        document.add_picture(corr_path, width=Inches(5.5))

    document.add_heading("Column Reference", level=2)
    col_rows = []
    for name, info in cols_info.items():
        col_rows.append({
            "column": name, "type": info["detected_type"],
            "missing_%": info.get("pct_missing"), "n_unique": info.get("n_unique"),
        })
    _add_list_of_dict_table(document, "", col_rows)
    document.add_page_break()

    # --- Agent narrative findings -------------------------------------------
    document.add_heading("Analysis Findings", level=1)
    if not report_sections:
        document.add_paragraph("The analysis agent did not record any narrative findings.")
    for section in report_sections:
        document.add_heading(section.title, level=2)
        document.add_paragraph(section.body)
        for fig_path in section.figures:
            if os.path.exists(fig_path):
                document.add_picture(fig_path, width=Inches(5.5))
    document.add_page_break()

    # --- Statistical results tables ----------------------------------------
    document.add_heading("Statistical Results", level=1)
    if not recorded_results:
        document.add_paragraph("No statistical tests were recorded via record_result() during this run.")
    for result in recorded_results:
        render_result_dict(document, result)

    # --- Validation & caveats ------------------------------------------------
    if validation_warnings:
        document.add_heading("Assumption Checks & Caveats", level=1)
        for w in validation_warnings:
            document.add_paragraph(w, style="List Bullet")

    # --- Ground truth comparison --------------------------------------------
    if ground_truth_comparison:
        document.add_heading("Comparison with Ground Truth", level=1)
        for key, info in ground_truth_comparison.items():
            document.add_heading(key, level=3)
            _add_kv_table(document, list(info.items()))

    document.add_page_break()
    document.add_heading("Methodology Notes", level=1)
    document.add_paragraph(
        "Statistical tests were selected and run by an LLM-based ReAct agent operating inside a "
        "restricted code sandbox, using a fixed library of statsmodels/scipy/pingouin wrapper "
        "functions (see stats_toolkit.py). Assumption checks (normality via Shapiro-Wilk, equal "
        "variance via Levene's test, homoscedasticity via Breusch-Pagan, etc.) are reported alongside "
        "each test result. Because the underlying language model is small and runs locally, treat "
        "this report as a fast first analysis pass -- verify any decision-critical conclusion "
        "independently before relying on it."
    )

    document.save(output_path)
    return output_path